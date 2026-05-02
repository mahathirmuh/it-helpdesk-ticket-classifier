from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title
title = doc.add_paragraph()
title_run = title.add_run("Bab 4. Metodologi")
title_run.font.size = Pt(16)
title_run.bold = True
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# 4.1 Dataset
doc.add_heading("4.1 Dataset dan Deskripsi Data", level=1)

doc.add_paragraph(
    "Penelitian ini menggunakan dataset tiket helpdesk IT bernama COBACEK yang dikumpulkan dari sistem "
    "ticketing internal. Dataset ini berisi 16.338 tiket helpdesk dengan informasi deskripsi masalah, "
    "kategori, dan prioritas yang telah dilabel oleh tim IT support."
)

table = doc.add_table(rows=5, cols=3)
table.style = 'Light Grid Accent 1'
header_cells = table.rows[0].cells
header_cells[0].text = "Atribut"
header_cells[1].text = "Deskripsi"
header_cells[2].text = "Nilai"

table.rows[1].cells[0].text = "Total Tiket"
table.rows[1].cells[1].text = "Jumlah sampel dalam dataset"
table.rows[1].cells[2].text = "16.338"

table.rows[2].cells[0].text = "Jumlah Kategori"
table.rows[2].cells[1].text = "Kategori IT yang unik"
table.rows[2].cells[2].text = "81"

table.rows[3].cells[0].text = "Prioritas"
table.rows[3].cells[1].text = "Level prioritas tiket"
table.rows[3].cells[2].text = "3 (Low, Medium, High)"

doc.add_paragraph()

doc.add_paragraph(
    "Setiap tiket memiliki tiga komponen data utama: (1) deskripsi masalah (description), "
    "(2) kategori IT (category), dan (3) level prioritas (priority). Dataset memiliki karakteristik "
    "imbalanced karena beberapa kategori hanya memiliki 1-5 sampel, sementara kategori lain memiliki "
    "puluhan hingga ratusan sampel. Distribusi kategori yang tidak merata ini mencerminkan pola real-world "
    "helpdesk tickets."
)

doc.add_paragraph()

# 4.2 Feature Extraction
doc.add_heading("4.2 Preprocessing dan Feature Extraction", level=1)

doc.add_paragraph(
    "Sebelum model dilatih, deskripsi tiket melalui tahap preprocessing standar: konversi ke lowercase, "
    "penghapusan whitespace berlebih, dan penghapusan karakter spesial. Feature extraction menggunakan "
    "TF-IDF (Term Frequency-Inverse Document Frequency) dari deskripsi tiket."
)

para_eq = doc.add_paragraph()
eq_run = para_eq.add_run("x_i = TF-IDF(d_i, ngram=(1,2))")
eq_run.italic = True
para_eq.add_run("                                         (1)")

doc.add_paragraph(
    "Dimana x_i adalah feature vector berdimensi ~1000+ untuk tiket ke-i, d_i adalah deskripsi tiket, "
    "dan ngram=(1,2) menggunakan unigram dan bigram. TF-IDF dipilih karena efektif untuk text classification "
    "dengan model linear seperti SVM, dan memberikan interpretasi yang jelas untuk fitur yang paling penting.",
    style="List Bullet"
)

doc.add_paragraph()

# 4.3 Experimental Setup
doc.add_heading("4.3 Experimental Setup", level=1)

doc.add_paragraph(
    "Penelitian ini membandingkan lima pendekatan klasifikasi: SVM (Support Vector Machine), "
    "Random Forest (RF), Logistic Regression (LR), BERT (Bidirectional Encoder Representations from Transformers), "
    "dan Hybrid SVM+GenAI. Untuk menjaga keadilan evaluasi dan menangani keterbatasan data, kami menggunakan "
    "Stratified K-Fold Cross-Validation dengan K=5 folds."
)

doc.add_paragraph()

doc.add_heading("4.3.1 Stratified K-Fold Cross-Validation", level=2)

doc.add_paragraph(
    "Stratified K-Fold dipilih karena dataset memiliki distribusi kategori yang imbalanced. "
    "Stratifikasi memastikan bahwa setiap fold mempertahankan proporsi kelas yang sama dengan dataset keseluruhan, "
    "sehingga hasil evaluasi lebih representative dan menghindari bias dari sampling yang tidak seimbang."
)

para_eq2 = doc.add_paragraph()
eq_run2 = para_eq2.add_run("Fold_k = {samples | class_proportion ≈ global_proportion}")
eq_run2.italic = True
para_eq2.add_run("                    (2)")

doc.add_paragraph(
    "Dengan K=5, setiap sampel diuji tepat satu kali (out-of-fold evaluation), dan hasil dari kelima fold "
    "digabungkan untuk menghitung metrik final. Pendekatan ini memberikan evaluasi yang stabil dan dapat direplikasi."
)

doc.add_paragraph()

doc.add_heading("4.3.2 Model Configurations", level=2)

doc.add_heading("Support Vector Machine (SVM)", level=3)
doc.add_paragraph(
    "SVM menggunakan kernel RBF (Radial Basis Function) untuk menangkap non-linear patterns. "
    "Dua model SVM terpisah dilatih: satu untuk klasifikasi kategori (81 kelas) dan satu untuk prioritas (3 kelas). "
    "Pipeline SVM: TF-IDF -> SVM classifier dengan default parameters scikit-learn."
)

doc.add_heading("Random Forest (RF)", level=3)
doc.add_paragraph(
    "Random Forest dilatih dengan n_estimators=200 trees dan random_state=42 untuk reproducibility. "
    "Seperti SVM, dua model RF terpisah untuk kategori dan prioritas. RF dipilih untuk membandingkan "
    "ensemble tree-based approach vs linear SVM."
)

doc.add_heading("Logistic Regression (LR)", level=3)
doc.add_paragraph(
    "Logistic Regression dengan max_iter=1000 sebagai baseline linear probabilistic model. "
    "LR lebih sederhana dari SVM tapi tetap powerful untuk text classification, sehingga memberikan "
    "good baseline untuk membandingkan model complexity vs performance."
)

doc.add_heading("BERT (Transformer-based)", level=3)
doc.add_paragraph(
    "Untuk deep learning baseline, kami menggunakan DistilBERT (distilbert-base-multilingual-cased) "
    "yang 40% lebih kecil dari BERT standard tapi tetap efektif untuk teks multibahasa. "
    "DistilBERT dilatih dengan 3 epochs, batch_size=16, dan max_token_length=128. "
    "Dua model DistilBERT terpisah untuk kategori dan prioritas."
)

doc.add_paragraph()

# 4.4 Metrics
doc.add_heading("4.4 Evaluation Metrics", level=1)

para_acc = doc.add_paragraph()
eq_acc = para_acc.add_run("Accuracy = (TP + TN) / (TP + TN + FP + FN)")
eq_acc.italic = True
para_acc.add_run("                        (3)")

doc.add_paragraph(
    "Mengukur proporsi prediksi yang benar dari total sampel. Accuracy memberikan gambaran overall "
    "performa model tetapi kurang sensitive terhadap class imbalance.",
    style="List Bullet"
)

para_prec = doc.add_paragraph()
eq_prec = para_prec.add_run("Macro Precision = (1/C) Σ TP_c / (TP_c + FP_c)")
eq_prec.italic = True
para_prec.add_run("                   (4)")

doc.add_paragraph(
    "Precision rata-rata untuk setiap kelas dengan bobot sama. Macro Precision lebih adil untuk dataset "
    "imbalanced karena tidak mengabaikan kelas-kelas minor.",
    style="List Bullet"
)

para_rec = doc.add_paragraph()
eq_rec = para_rec.add_run("Macro Recall = (1/C) Σ TP_c / (TP_c + FN_c)")
eq_rec.italic = True
para_rec.add_run("                      (5)")

doc.add_paragraph("Recall rata-rata untuk setiap kelas dengan bobot sama.", style="List Bullet")

para_f1 = doc.add_paragraph()
eq_f1 = para_f1.add_run("Macro F1 = (1/C) Σ 2·P_c·R_c / (P_c + R_c)")
eq_f1.italic = True
para_f1.add_run("                    (6)")

doc.add_paragraph(
    "F1-score harmonic mean antara precision dan recall untuk setiap kelas, dengan bobot sama. "
    "Macro F1 adalah metrik utama penelitian ini karena balance antara precision dan recall, "
    "dan fair terhadap class imbalance.",
    style="List Bullet"
)

para_wf1 = doc.add_paragraph()
eq_wf1 = para_wf1.add_run("Weighted F1 = Σ (count_c / C) × F1_c")
eq_wf1.italic = True
para_wf1.add_run("                        (7)")

doc.add_paragraph(
    "F1-score weighted berdasarkan frekuensi setiap kelas. Weighted F1 lebih sensitive terhadap "
    "majority classes dan berguna untuk melihat impact dari class imbalance.",
    style="List Bullet"
)

doc.add_paragraph()

# 4.5 Hybrid SVM
doc.add_heading("4.5 Hybrid SVM + GenAI Architecture", level=1)

doc.add_paragraph(
    "Pendekatan hybrid menggabungkan SVM baseline dengan generative AI (OpenAI API) untuk koreksi prediksi "
    "yang salah. Pipeline terdiri dari empat tahap:"
)

doc.add_heading("Tahap 1: SVM Baseline Prediction", level=3)
doc.add_paragraph("Model SVM melakukan prediksi kategori dan prioritas untuk semua 16.338 tiket.")

doc.add_heading("Tahap 2: Mismatch Detection", level=3)

para_mismatch = doc.add_paragraph()
eq_mismatch = para_mismatch.add_run("M_i = {i | (ŷ^{cat}_i ≠ c_i) ∨ (ŷ^{pri}_i ≠ p_i)}")
eq_mismatch.italic = True
para_mismatch.add_run("            (8)")

doc.add_paragraph(
    "Dimana M_i adalah set indeks tiket yang salah prediksi (mismatch). Tiket dalam M_i adalah kandidat "
    "untuk koreksi oleh GenAI karena SVM prediksi mereka tidak match dengan ground truth."
)

doc.add_heading("Tahap 3: GenAI Correction", level=3)

para_genai = doc.add_paragraph()
eq_genai = para_genai.add_run("(ŷ^{cat,hybrid}_i, ŷ^{pri,hybrid}_i) = GenAI(d_i, ŷ^{cat}_i, ŷ^{pri}_i)  untuk i ∈ M_i")
eq_genai.italic = True
para_genai.add_run("  (9)")

doc.add_paragraph(
    "Untuk setiap tiket dalam M_i, GenAI (menggunakan OpenAI API) menerima deskripsi tiket d_i dan "
    'prediksi SVM, kemudian mengembalikan kategori dan prioritas yang diperbaiki. GenAI diminta untuk '
    '"validate and refine" prediksi SVM hanya jika diperlukan, untuk menghindari over-correction.'
)

doc.add_heading("Tahap 4: Hybrid Result Combination", level=3)

para_hybrid = doc.add_paragraph()
eq_hybrid = para_hybrid.add_run("Hybrid = {SVM correct rows} ∪ {GenAI corrected rows}")
eq_hybrid.italic = True
para_hybrid.add_run("           (10)")

doc.add_paragraph(
    "Hasil akhir hybrid adalah gabungan antara: (1) baris yang SVM sudah benar (tidak perlu koreksi), "
    "dan (2) baris yang telah dikoreksi oleh GenAI. Hasil hybrid ini dievaluasi dengan metrik yang sama "
    "seperti SVM, RF, LR, dan BERT untuk fair comparison."
)

doc.add_paragraph()

# 4.6 Implementation
doc.add_heading("4.6 Implementation dan Tools", level=1)

doc.add_paragraph(
    "Seluruh eksperimen diimplementasikan menggunakan Python dengan libraries berikut: "
    "scikit-learn untuk SVM/RF/LR, PyTorch dan Transformers untuk BERT, dan OpenAI API untuk GenAI. "
    "Cross-validation dan evaluasi metrik menggunakan scikit-learn. Hasil eksperimen disimpan dalam "
    "format Excel (cobacek_compare.xlsx) dengan sheet terpisah untuk predictions, metrics, dan summary."
)

from pathlib import Path
import os
ROOT = Path(__file__).resolve().parents[2]
os.chdir(ROOT)
doc.save("paper/Bab_4_Metodologi.docx")
print("[OK] Document created: paper/Bab_4_Metodologi.docx")
