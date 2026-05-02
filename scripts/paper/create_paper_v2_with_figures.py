import sys, os
sys.stdout.reconfigure(encoding="utf-8")

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import numpy as np
import pandas as pd
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Resolve paths relative to project root (script is at scripts/paper/)
ROOT       = Path(__file__).resolve().parents[2]
DATA_FILE  = ROOT / "data"    / "cobacek.xlsx"
RESULTS    = ROOT / "results" / "cobacek_compare_final.xlsx"
FIG_DIR    = ROOT / "paper"   / "figures"
PAPER_OUT  = ROOT / "paper"   / "IT_Helpdesk_Ticket_Classifier_Paper_V2.docx"
FIG_DIR.mkdir(parents=True, exist_ok=True)
os.chdir(ROOT)  # so relative paths in figs work consistently

# ======================================================================
# HELPERS
# ======================================================================

BLUE   = "#1565C0"
LBLUE  = "#42A5F5"
GREEN  = "#2E7D32"
ORANGE = "#E65100"
RED    = "#C62828"
GREY   = "#546E7A"
WHITE  = "#FFFFFF"

def save_fig(name):
    plt.savefig(name, dpi=150, bbox_inches="tight", facecolor="white")
    plt.close()

def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(10)

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)
    return h

def add_bold_para(doc, parts):
    """parts = list of (text, bold)"""
    p = doc.add_paragraph()
    for text, bold in parts:
        r = p.add_run(text)
        r.bold = bold
        r.font.size = Pt(11)
    return p

def add_para(doc, text, style=None):
    p = doc.add_paragraph(text, style=style) if style else doc.add_paragraph(text)
    for run in p.runs:
        run.font.size = Pt(11)
    return p


# ======================================================================
# FIGURE 1 — Category Distribution (top 20)
# ======================================================================
def make_fig1_category():
    df = pd.read_excel(DATA_FILE)
    cats = df["category"].value_counts().head(20)

    fig, ax = plt.subplots(figsize=(10, 7))
    colors = ["#1565C0" if v >= 1000 else "#42A5F5" if v >= 100 else "#BBDEFB"
              for v in cats.values]
    bars = ax.barh(range(len(cats)), cats.values, color=colors, height=0.7,
                   edgecolor="white", linewidth=0.5)
    ax.set_yticks(range(len(cats)))
    ax.set_yticklabels(cats.index, fontsize=10)
    ax.invert_yaxis()
    ax.set_xlabel("Jumlah Tiket", fontsize=11)
    ax.set_title(
        "Distribusi 20 Kategori Teratas (dari 81 Total)\nDataset COBACEK  (n = 16.338)",
        fontsize=13, fontweight="bold", pad=12
    )
    for bar, val in zip(bars, cats.values):
        ax.text(bar.get_width() + 30, bar.get_y() + bar.get_height() / 2,
                f"{val:,}", va="center", fontsize=9, color="#333")
    ax.set_xlim(0, cats.values.max() * 1.18)
    ax.grid(axis="x", alpha=0.25, linestyle="--")
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)

    # legend
    handles = [
        mpatches.Patch(color="#1565C0", label=">= 1.000 sampel"),
        mpatches.Patch(color="#42A5F5", label="100 – 999 sampel"),
        mpatches.Patch(color="#BBDEFB", label="< 100 sampel"),
    ]
    ax.legend(handles=handles, loc="lower right", fontsize=9)
    save_fig("paper/figures/fig1_category_dist.png")
    print("[OK] fig1_category_dist.png")


# ======================================================================
# FIGURE 2 — Priority Distribution
# ======================================================================
def make_fig2_priority():
    df = pd.read_excel(DATA_FILE)
    order  = ["high", "medium", "low"]
    labels = ["High", "Medium", "Low"]
    prio   = df["priority"].value_counts().reindex(order)

    fig, ax = plt.subplots(figsize=(6, 4))
    colors = ["#C62828", "#EF6C00", "#2E7D32"]
    bars = ax.bar(labels, prio.values, color=colors, width=0.5, edgecolor="white")
    for bar, val in zip(bars, prio.values):
        pct = val / prio.sum() * 100
        ax.text(bar.get_x() + bar.get_width() / 2,
                bar.get_height() + 100,
                f"{val:,}\n({pct:.1f}%)",
                ha="center", fontsize=11, fontweight="bold")
    ax.set_ylabel("Jumlah Tiket", fontsize=11)
    ax.set_title("Distribusi Level Prioritas  (n = 16.338)", fontsize=13, fontweight="bold")
    ax.set_ylim(0, prio.max() * 1.25)
    ax.grid(axis="y", alpha=0.25, linestyle="--")
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    save_fig("paper/figures/fig2_priority_dist.png")
    print("[OK] fig2_priority_dist.png")


# ======================================================================
# FIGURE 3 — Research Pipeline (horizontal flowchart)
# ======================================================================
def make_fig3_pipeline():
    fig, ax = plt.subplots(figsize=(13, 3.5))
    ax.set_xlim(0, 13)
    ax.set_ylim(0, 3.5)
    ax.axis("off")
    ax.set_facecolor("white")

    def fbox(cx, cy, w, h, lines, fc=BLUE, fs=9):
        r = mpatches.FancyBboxPatch(
            (cx - w/2, cy - h/2), w, h,
            boxstyle="round,pad=0.12", facecolor=fc, edgecolor="none", zorder=3
        )
        ax.add_patch(r)
        ax.text(cx, cy, "\n".join(lines), ha="center", va="center",
                fontsize=fs, color=WHITE, fontweight="bold",
                multialignment="center", zorder=4, linespacing=1.4)

    def arr(x1, y, x2, color=GREY):
        ax.annotate("", xy=(x2, y), xytext=(x1, y),
                    arrowprops=dict(arrowstyle="->", color=color, lw=1.8), zorder=2)

    # boxes
    fbox(1.2, 1.75, 1.8, 1.4, ["Dataset", "COBACEK", "(16.338 tikets)"],         fc="#1565C0")
    fbox(3.3, 1.75, 1.8, 1.4, ["Text", "Preprocessing", "(clean, lower)"],        fc="#1976D2")
    fbox(5.4, 1.75, 1.8, 1.4, ["Feature", "Extraction", "(TF-IDF, ngram 1-2)"],   fc="#1E88E5")
    fbox(7.5, 2.75, 1.8, 0.9, ["SVM  (Linear)"],                                   fc=GREEN)
    fbox(7.5, 1.75, 1.8, 0.9, ["RF / LR"],                                         fc=GREEN)
    fbox(7.5, 0.75, 1.8, 0.9, ["Hybrid SVM+GenAI"],                                fc=ORANGE)
    fbox(10.7, 1.75, 2.0, 1.4, ["Evaluasi", "5-Fold Stratified", "Accuracy, F1"], fc="#6A1B9A")

    # arrows main flow
    arr(2.1, 1.75, 2.4)
    arr(4.2, 1.75, 4.5)
    # branch from feature extraction to models
    ax.annotate("", xy=(6.6, 2.75), xytext=(6.3, 1.75),
                arrowprops=dict(arrowstyle="->", color=GREY, lw=1.5))
    arr(6.3, 1.75, 6.6)
    ax.annotate("", xy=(6.6, 0.75), xytext=(6.3, 1.75),
                arrowprops=dict(arrowstyle="->", color=GREY, lw=1.5))
    # models to eval
    ax.annotate("", xy=(9.7, 1.75), xytext=(8.4, 2.75),
                arrowprops=dict(arrowstyle="->", color=GREY, lw=1.5))
    arr(8.4, 1.75, 9.7)
    ax.annotate("", xy=(9.7, 1.75), xytext=(8.4, 0.75),
                arrowprops=dict(arrowstyle="->", color=GREY, lw=1.5))

    ax.set_title("Gambar 3. Alur Metodologi Penelitian", fontsize=12, fontweight="bold", pad=6)
    save_fig("paper/figures/fig3_pipeline.png")
    print("[OK] fig3_pipeline.png")


# ======================================================================
# FIGURE 4 — Hybrid SVM+GenAI Architecture (vertical flowchart)
# ======================================================================
def make_fig4_hybrid():
    fig, ax = plt.subplots(figsize=(8, 9))
    ax.set_xlim(0, 8)
    ax.set_ylim(0, 9)
    ax.axis("off")

    def fbox(cx, cy, w, h, lines, fc=BLUE, fs=9.5):
        r = mpatches.FancyBboxPatch(
            (cx - w/2, cy - h/2), w, h,
            boxstyle="round,pad=0.15", facecolor=fc, edgecolor="white",
            linewidth=1.5, zorder=3
        )
        ax.add_patch(r)
        ax.text(cx, cy, "\n".join(lines), ha="center", va="center",
                fontsize=fs, color=WHITE, fontweight="bold",
                multialignment="center", zorder=4, linespacing=1.4)

    def diamond(cx, cy, w, h, text, fc="#F57F17"):
        xs = [cx, cx+w/2, cx, cx-w/2, cx]
        ys = [cy+h/2, cy, cy-h/2, cy, cy+h/2]
        ax.fill(xs, ys, facecolor=fc, edgecolor="white", linewidth=1.5, zorder=3)
        ax.text(cx, cy, text, ha="center", va="center", fontsize=9.5,
                color=WHITE, fontweight="bold", multialignment="center", zorder=4)

    def arrow_v(cx, y1, y2, color=GREY, label=""):
        ax.annotate("", xy=(cx, y2), xytext=(cx, y1),
                    arrowprops=dict(arrowstyle="->", color=color, lw=1.8), zorder=2)
        if label:
            ax.text(cx + 0.15, (y1+y2)/2, label, fontsize=8.5, color=color, va="center")

    def arrow_h(x1, cy, x2, color=GREY, label=""):
        ax.annotate("", xy=(x2, cy), xytext=(x1, cy),
                    arrowprops=dict(arrowstyle="->", color=color, lw=1.8), zorder=2)
        if label:
            ax.text((x1+x2)/2, cy + 0.15, label, fontsize=8.5, color=color, ha="center")

    # Stage 1
    fbox(4, 8.3, 4.5, 0.9,
         ["Tahap 1: SVM Baseline Prediction",
          "16.338 tikets -> prediksi kategori & prioritas"],
         fc=GREEN, fs=9)
    arrow_v(4, 7.85, 7.35)

    # Stage 2 — diamond
    diamond(4, 7.0, 4.5, 0.7, "Tahap 2: Mismatch Detection\nSalah kategori ATAU salah prioritas?")
    ax.text(5.7, 7.0, "Tidak\n(9.731 baris, 59.6%)", fontsize=8.5, color=GREEN, va="center", ha="left")
    ax.text(4.15, 6.45, "Ya\n(6.607 baris, 40.4%)", fontsize=8.5, color=ORANGE, va="top", ha="center")

    # Branch YES (down) and NO (right)
    arrow_v(4, 6.65, 6.15, color=ORANGE)           # mismatch -> GenAI
    arrow_h(5.7, 7.0, 7.2, color=GREEN)             # correct -> right merge
    # merge box
    fbox(7.2, 5.5, 1.2, 2.8, ["Prediksi\nSVM\nsudah\nbenar"], fc=GREEN, fs=8.5)

    # Stage 3
    fbox(4, 5.5, 3.5, 0.9,
         ["Tahap 3: GenAI Correction",
          "OpenAI gpt-5.4-mini mengoreksi\nkategori & prioritas tiket mismatch"],
         fc=ORANGE, fs=8.5)
    arrow_v(4, 5.05, 4.35)

    # Stage 4
    fbox(4, 3.9, 3.5, 0.75,
         ["Tahap 4: Hybrid Result Combination",
          "SVM correct rows + GenAI corrected rows"],
         fc=BLUE, fs=8.5)
    # merge arrows
    ax.annotate("", xy=(5.75, 3.9), xytext=(6.6, 4.5),
                arrowprops=dict(arrowstyle="->", color=GREEN, lw=1.5), zorder=2)
    arrow_v(4, 3.525, 3.05)

    # Evaluation
    fbox(4, 2.6, 3.5, 0.75,
         ["Evaluasi vs SVM / RF / LR",
          "Accuracy, Macro F1, Weighted F1"],
         fc="#6A1B9A", fs=8.5)

    # Result annotation
    ax.text(4, 1.9,
            "Hasil: Category Acc turun 14.05 pp (81.22% -> 66.17%)\n"
            "Priority Acc naik 1.33 pp (72.50% -> 73.83%)  |  Waktu: 144x lebih lambat",
            ha="center", va="center", fontsize=9, color=RED,
            fontweight="bold", style="italic",
            bbox=dict(facecolor="#FFEBEE", edgecolor=RED, boxstyle="round,pad=0.3"))

    ax.set_title("Gambar 4. Arsitektur Pipeline Hybrid SVM + GenAI",
                 fontsize=12, fontweight="bold", pad=8)
    save_fig("paper/figures/fig4_hybrid_arch.png")
    print("[OK] fig4_hybrid_arch.png")


# ======================================================================
# FIGURE 5 — Model Accuracy Comparison
# ======================================================================
def make_fig5_accuracy():
    metrics = pd.read_excel(RESULTS, sheet_name="Metrics")
    cat = metrics[metrics["label"] == "category"].set_index("approach")["accuracy"]
    pri = metrics[metrics["label"] == "priority"].set_index("approach")["accuracy"]

    models   = ["SVM", "Random Forest", "Logistic Regression", "Hybrid SVM\n(gpt-5.4-mini)"]
    keys_cat = ["SVM", "Random Forest", "Logistic Regression", "Hybrid SVM (gpt-5.4-mini)"]
    cat_vals = [cat[k] for k in keys_cat]
    pri_vals = [pri[k] for k in keys_cat]

    x = np.arange(len(models))
    w = 0.35

    fig, ax = plt.subplots(figsize=(10, 5.5))
    b1 = ax.bar(x - w/2, cat_vals, w, label="Kategori (81 kelas)", color=LBLUE,  edgecolor="white")
    b2 = ax.bar(x + w/2, pri_vals, w, label="Prioritas (3 kelas)",  color=ORANGE, edgecolor="white")

    for b, v in zip(b1, cat_vals):
        ax.text(b.get_x() + b.get_width()/2, v + 0.006, f"{v:.4f}",
                ha="center", va="bottom", fontsize=9, fontweight="bold")
    for b, v in zip(b2, pri_vals):
        ax.text(b.get_x() + b.get_width()/2, v + 0.006, f"{v:.4f}",
                ha="center", va="bottom", fontsize=9, fontweight="bold")

    # Annotate accuracy paradox arrow
    hybrid_idx = 3
    ax.annotate(
        "Turun 14.05 pp!\n(Accuracy Paradox)",
        xy=(x[hybrid_idx] - w/2, cat_vals[hybrid_idx] + 0.03),
        xytext=(x[hybrid_idx] - w/2 - 0.6, cat_vals[hybrid_idx] + 0.1),
        arrowprops=dict(arrowstyle="->", color=RED, lw=1.5),
        color=RED, fontsize=9, fontweight="bold"
    )

    ax.set_xticks(x)
    ax.set_xticklabels(models, fontsize=10)
    ax.set_ylabel("Accuracy", fontsize=11)
    ax.set_ylim(0, 1.0)
    ax.set_title("Perbandingan Accuracy Antar Model\n(Dataset COBACEK, n = 16.338)",
                 fontsize=13, fontweight="bold")
    ax.legend(fontsize=10)
    ax.yaxis.set_major_formatter(plt.FuncFormatter(lambda v, _: f"{v:.0%}"))
    ax.grid(axis="y", alpha=0.3, linestyle="--")
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    save_fig("paper/figures/fig5_accuracy_compare.png")
    print("[OK] fig5_accuracy_compare.png")


# ======================================================================
# FIGURE 6 — Macro F1 Comparison
# ======================================================================
def make_fig6_macro_f1():
    metrics = pd.read_excel(RESULTS, sheet_name="Metrics")
    cat = metrics[metrics["label"] == "category"].set_index("approach")["macro_f1"]
    pri = metrics[metrics["label"] == "priority"].set_index("approach")["macro_f1"]

    models   = ["SVM", "Random Forest", "Logistic Regression", "Hybrid SVM\n(gpt-5.4-mini)"]
    keys_cat = ["SVM", "Random Forest", "Logistic Regression", "Hybrid SVM (gpt-5.4-mini)"]
    cat_vals = [cat[k] for k in keys_cat]
    pri_vals = [pri[k] for k in keys_cat]

    x = np.arange(len(models))
    w = 0.35

    fig, ax = plt.subplots(figsize=(10, 5.5))
    b1 = ax.bar(x - w/2, cat_vals, w, label="Kategori (81 kelas)", color="#1E88E5", edgecolor="white")
    b2 = ax.bar(x + w/2, pri_vals, w, label="Prioritas (3 kelas)",  color="#FB8C00", edgecolor="white")

    for b, v in zip(b1, cat_vals):
        ax.text(b.get_x() + b.get_width()/2, v + 0.004, f"{v:.4f}",
                ha="center", va="bottom", fontsize=9, fontweight="bold")
    for b, v in zip(b2, pri_vals):
        ax.text(b.get_x() + b.get_width()/2, v + 0.004, f"{v:.4f}",
                ha="center", va="bottom", fontsize=9, fontweight="bold")

    ax.set_xticks(x)
    ax.set_xticklabels(models, fontsize=10)
    ax.set_ylabel("Macro F1-Score", fontsize=11)
    ax.set_ylim(0, 1.0)
    ax.set_title("Perbandingan Macro F1-Score Antar Model\n(Macro F1 rendah untuk kategori mencerminkan extreme class imbalance)",
                 fontsize=12, fontweight="bold")
    ax.legend(fontsize=10)
    ax.grid(axis="y", alpha=0.3, linestyle="--")
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)

    # annotation about low macro f1
    ax.text(0.01, 0.95,
            "Catatan: Macro F1 kategori rendah (0.11-0.29)\nbukan berarti model buruk,\n"
            "tapi karena banyak kategori dengan < 5 sampel.",
            transform=ax.transAxes, fontsize=8.5,
            verticalalignment="top",
            bbox=dict(boxstyle="round", facecolor="#E3F2FD", alpha=0.7))
    save_fig("paper/figures/fig6_macro_f1_compare.png")
    print("[OK] fig6_macro_f1_compare.png")


# ======================================================================
# FIGURE 7 — Computational Time (log scale)
# ======================================================================
def make_fig7_time():
    metrics = pd.read_excel(RESULTS, sheet_name="Metrics")
    times = (metrics[metrics["label"] == "category"]
             .set_index("approach")["elapsed_seconds"])

    models = ["SVM", "Random Forest", "Logistic\nRegression", "Hybrid SVM\n(gpt-5.4-mini)"]
    keys   = ["SVM", "Random Forest", "Logistic Regression", "Hybrid SVM (gpt-5.4-mini)"]
    vals   = [times[k] for k in keys]
    colors = [GREEN, "#43A047", "#66BB6A", RED]

    fig, ax = plt.subplots(figsize=(8, 5))
    bars = ax.bar(models, vals, color=colors, width=0.5, edgecolor="white")
    ax.set_yscale("log")
    ax.set_ylabel("Waktu (detik, skala log)", fontsize=11)
    ax.set_title("Perbandingan Waktu Komputasi Antar Model\n(skala logaritmik)",
                 fontsize=13, fontweight="bold")

    for bar, val in zip(bars, vals):
        mins = val / 60
        label = f"{val:.1f} s\n({mins:.1f} mnt)" if val >= 60 else f"{val:.1f} s"
        ax.text(bar.get_x() + bar.get_width()/2,
                bar.get_height() * 1.3,
                label, ha="center", va="bottom", fontsize=9.5, fontweight="bold")

    # highlight hybrid
    ax.annotate(
        "144x lebih lambat\ndari SVM",
        xy=(3, vals[3]),
        xytext=(2.3, vals[3] * 0.35),
        arrowprops=dict(arrowstyle="->", color=RED, lw=1.5),
        color=RED, fontsize=9.5, fontweight="bold"
    )

    ax.grid(axis="y", alpha=0.25, linestyle="--", which="both")
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.set_ylim(1, vals[3] * 5)
    save_fig("paper/figures/fig7_time_compare.png")
    print("[OK] fig7_time_compare.png")


# ======================================================================
# GENERATE ALL FIGURES
# ======================================================================
print("Generating figures...")
make_fig1_category()
make_fig2_priority()
make_fig3_pipeline()
make_fig4_hybrid()
make_fig5_accuracy()
make_fig6_macro_f1()
make_fig7_time()
print("All figures generated.\n")


# ======================================================================
# READ ACTUAL METRICS
# ======================================================================
metrics_df = pd.read_excel(RESULTS, sheet_name="Metrics")
summary_df = pd.read_excel(RESULTS, sheet_name="Summary")

def get_metric(approach, label, metric):
    row = metrics_df[(metrics_df["approach"] == approach) & (metrics_df["label"] == label)]
    return row[metric].values[0] if len(row) else "N/A"


# ======================================================================
# BUILD WORD DOCUMENT
# ======================================================================
doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

# ---- TITLE ----
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = title_p.add_run(
    "IT Helpdesk Ticket Classifier:\n"
    "Perbandingan SVM, Random Forest, Logistic Regression, dan Hybrid SVM+GenAI"
)
tr.bold = True
tr.font.size = Pt(16)
tr.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)

doc.add_paragraph()
auth = doc.add_paragraph()
auth.alignment = WD_ALIGN_PARAGRAPH.CENTER
auth.add_run("Mahathir Muhammad").font.size = Pt(12)

yr = doc.add_paragraph()
yr.alignment = WD_ALIGN_PARAGRAPH.CENTER
yr.add_run("2026").font.size = Pt(11)

doc.add_paragraph()

# ---- ABSTRACT ----
abs_head = doc.add_paragraph()
abs_run = abs_head.add_run("Abstrak")
abs_run.bold = True
abs_run.font.size = Pt(12)
abs_run.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)
abs_head.alignment = WD_ALIGN_PARAGRAPH.CENTER

abs_body = doc.add_paragraph(
    "Penelitian ini membandingkan empat pendekatan klasifikasi teks untuk otomasi sistem IT helpdesk ticket "
    "classification: Support Vector Machine (SVM) dengan Linear kernel, Random Forest (RF), Logistic Regression "
    "(LR), dan Hybrid SVM+GenAI. Dataset yang digunakan adalah COBACEK dengan 16.338 tiket helpdesk, 81 kategori, "
    "dan 3 level prioritas (High, Medium, Low). Eksperimen menggunakan Stratified 5-Fold Cross-Validation dengan "
    "TF-IDF (ngram 1-2) sebagai feature extraction. Hasil menunjukkan SVM mencapai akurasi tertinggi untuk "
    "klasifikasi kategori (81.22%, Macro F1 = 0.2587) dan kompetitif untuk prioritas (72.50%). Temuan utama "
    "adalah discovery 'accuracy paradox' pada Hybrid SVM+GenAI: meskipun 6.607 prediksi (40.4% dataset) "
    "dikoreksi oleh OpenAI gpt-5.4-mini, akurasi kategori justru turun 14.05 percentage points menjadi 66.17%. "
    "Hybrid hanya meningkatkan akurasi prioritas sebesar 1.33 pp dengan waktu komputasi 144x lebih lambat "
    "(9.772 detik vs 67 detik). Penelitian ini menunjukkan bahwa naive hybrid approach combining classical ML "
    "dengan GenAI tidak selalu meningkatkan performa, dan analisis mendalam diberikan tentang penyebab "
    "accuracy paradox tersebut."
)
for run in abs_body.runs:
    run.font.size = Pt(10.5)

kw = doc.add_paragraph()
kw_r = kw.add_run("Kata Kunci: ")
kw_r.bold = True
kw_r.font.size = Pt(10.5)
kw_r2 = kw.add_run(
    "klasifikasi teks, SVM, generative AI, hybrid model, IT helpdesk, class imbalance, accuracy paradox"
)
kw_r2.font.size = Pt(10.5)

doc.add_page_break()

# ===========================================================
# BAB 1 — PENDAHULUAN
# ===========================================================
add_heading(doc, "Bab 1. Pendahuluan", level=1)

add_heading(doc, "1.1 Latar Belakang", level=2)
add_para(doc,
    "Sistem ticketing helpdesk IT merupakan backbone operasional support dalam organisasi modern. "
    "Setiap hari, ribuan tiket masuk dengan variasi kategori dan tingkat prioritas yang berbeda-beda. "
    "Klasifikasi manual tiket oleh tim support memakan waktu, rentan error, dan tidak scalable seiring "
    "meningkatnya volume tiket. Otomasi klasifikasi tiket menggunakan machine learning menjadi solusi "
    "strategis untuk meningkatkan efisiensi operasional, mengurangi response time, dan memastikan tiket "
    "diarahkan ke tim yang tepat dengan prioritas yang akurat."
)
add_para(doc,
    "Penelitian ini mengeksplorasi beberapa pendekatan machine learning — dari classical ML seperti SVM, "
    "Random Forest, dan Logistic Regression, hingga hybrid approach yang menggabungkan SVM dengan Generative AI "
    "(OpenAI API). Pendekatan hybrid ini menarik secara teoritis karena berpotensi menggabungkan kecepatan "
    "classical ML dengan kemampuan reasoning GenAI untuk mengoreksi prediksi yang salah."
)

add_heading(doc, "1.2 Rumusan Masalah", level=2)
add_para(doc,
    "Penelitian ini menjawab pertanyaan: (1) Bagaimana performa relatif SVM, RF, LR untuk klasifikasi "
    "tiket helpdesk IT pada dataset imbalanced dengan 81 kategori? (2) Apakah hybrid SVM+GenAI dapat "
    "meningkatkan akurasi klasifikasi? (3) Apa trade-off antara akurasi, waktu komputasi, dan biaya "
    "operasional dari setiap pendekatan?"
)

add_heading(doc, "1.3 Tujuan Penelitian", level=2)
add_para(doc, "Penelitian ini bertujuan untuk:")
for item in [
    "Membandingkan performa SVM, RF, LR, dan Hybrid SVM+GenAI pada dataset real-world tiket helpdesk IT",
    "Menganalisis trade-off antara akurasi, scalability, dan biaya komputasi",
    "Mengevaluasi efektivitas hybrid approach (SVM + GenAI correction) untuk koreksi prediksi",
    "Mengidentifikasi best-practice untuk klasifikasi multi-label pada dataset imbalanced",
]:
    add_para(doc, item, style="List Number")

add_heading(doc, "1.4 Kontribusi Penelitian", level=2)
add_para(doc,
    "Kontribusi utama penelitian ini: (1) Comprehensive benchmark empat model pada dataset real helpdesk "
    "dengan 81 kategori; (2) Discovery dan analisis mendalam 'accuracy paradox' pada hybrid SVM+GenAI; "
    "(3) Insights tentang tantangan multi-label simultaneous correction oleh GenAI; "
    "(4) Practical recommendations untuk deployment di production helpdesk systems."
)

doc.add_page_break()

# ===========================================================
# BAB 2 — TINJAUAN LITERATUR
# ===========================================================
add_heading(doc, "Bab 2. Tinjauan Literatur", level=1)

add_heading(doc, "2.1 Machine Learning untuk Text Classification", level=2)
add_para(doc,
    "Text classification adalah task fundamental dalam natural language processing (NLP). Pipeline "
    "umumnya terdiri dari: (1) text preprocessing, (2) feature extraction, dan (3) classification model. "
    "Pendekatan klasik menggunakan Bag-of-Words (BoW) atau Term Frequency-Inverse Document Frequency (TF-IDF) "
    "untuk feature representation, kemudian diklasifikasi dengan SVM, Naive Bayes, atau Logistic Regression "
    "(Sebastiani, 2002; Yang & Liu, 1999)."
)
add_para(doc,
    "Support Vector Machine (SVM) telah terbukti sangat efektif untuk text classification karena kemampuannya "
    "bekerja di high-dimensional sparse feature spaces seperti TF-IDF (Joachims, 1998). Linear SVM khususnya "
    "efisien secara komputasi dan memberikan generalization yang baik pada dataset teks besar."
)
add_para(doc,
    "Dalam dekade terakhir, deep learning seperti BERT (Devlin et al., 2019) mendominasi NLP dengan "
    "menggunakan pre-trained transformer yang di-fine-tune untuk berbagai downstream tasks. Namun untuk "
    "dataset kecil atau medium dengan many-class classification, classical ML dengan TF-IDF seringkali "
    "tetap kompetitif dan lebih efisien secara komputasi (Sun et al., 2019)."
)

add_heading(doc, "2.2 Handling Class Imbalance", level=2)
add_para(doc,
    "Class imbalance adalah masalah umum dalam real-world classification. Dataset dengan distribusi kelas "
    "tidak merata menyebabkan bias terhadap majority classes dan poor generalization pada minority classes. "
    "Strategi penanganan meliputi: oversampling (SMOTE), undersampling, cost-sensitive learning, dan "
    "stratified sampling. Metrik evaluasi seperti Macro F1 lebih adil dibanding accuracy untuk "
    "imbalanced datasets karena memberikan bobot equal pada setiap kelas (He & Garcia, 2009)."
)
add_para(doc,
    "Dalam konteks IT helpdesk, class imbalance sangat umum — kategori seperti 'Security' dan 'Bug' "
    "mendominasi, sementara banyak kategori spesifik hanya memiliki 1-5 sampel. Hal ini menciptakan "
    "tantangan khusus dalam evaluasi: accuracy tinggi tidak selalu berarti model baik untuk minority classes."
)

add_heading(doc, "2.3 Hybrid Approaches dan GenAI Integration", level=2)
add_para(doc,
    "Hybrid approaches menggabungkan multiple models untuk meningkatkan robustness. Generative AI seperti "
    "GPT-4 (OpenAI, 2023) menunjukkan kemampuan luar biasa dalam text understanding dan reasoning. "
    "Beberapa penelitian mengeksplorasi kombinasi classical ML dengan LLM untuk meningkatkan akurasi, "
    "misalnya menggunakan LLM untuk data augmentation atau error correction (Wei et al., 2022)."
)
add_para(doc,
    "Namun, tantangan utama dalam hybrid approach adalah interaction effects antara model — koreksi GenAI "
    "tidak selalu meningkatkan akurasi dan dapat menimbulkan over-correction pada prediksi yang sudah benar. "
    "Confidence-based routing (hanya apply GenAI untuk low-confidence predictions) adalah pendekatan yang "
    "lebih targeted dan berpotensi lebih efektif (Ratner et al., 2019)."
)

doc.add_page_break()

# ===========================================================
# BAB 3 — METODOLOGI
# ===========================================================
add_heading(doc, "Bab 3. Metodologi", level=1)

add_heading(doc, "3.1 Dataset dan Deskripsi Data", level=2)
add_para(doc,
    "Penelitian ini menggunakan dataset tiket helpdesk IT bernama COBACEK yang dikumpulkan dari sistem "
    "ticketing internal. Dataset berisi 16.338 tiket dengan deskripsi masalah (description), "
    "kategori IT (category), dan level prioritas (priority) yang telah dilabel oleh tim IT support."
)

# Table dataset
tbl = doc.add_table(rows=5, cols=3)
tbl.style = "Light Grid Accent 1"
hdr = tbl.rows[0].cells
hdr[0].text, hdr[1].text, hdr[2].text = "Atribut", "Deskripsi", "Nilai"
rows_data = [
    ("Total Tiket",       "Jumlah sampel dalam dataset",  "16.338"),
    ("Jumlah Kategori",   "Kategori IT yang unik",         "81"),
    ("Level Prioritas",   "Skala prioritas tiket",         "3 (High, Medium, Low)"),
    ("Distribusi Tiket",  "Rasio majority/minority class", "Security 3.333 / kategori minor 1-5"),
]
for i, (a, b, c) in enumerate(rows_data, 1):
    tbl.rows[i].cells[0].text = a
    tbl.rows[i].cells[1].text = b
    tbl.rows[i].cells[2].text = c

doc.add_paragraph()
add_para(doc,
    "Dataset memiliki karakteristik extreme class imbalance: lima kategori teratas (Security, Bug, Feedback, "
    "Feature, Performance) menyumbang lebih dari 74% total tiket, sementara 50+ kategori lainnya hanya "
    "memiliki kurang dari 30 sampel masing-masing. Gambar 1 dan 2 di bawah memvisualisasikan distribusi ini."
)

# Figure 1
doc.add_picture("paper/figures/fig1_category_dist.png", width=Inches(5.8))
add_caption(doc, "Gambar 1. Distribusi 20 Kategori Teratas Dataset COBACEK (menunjukkan extreme class imbalance)")
doc.add_paragraph()

# Figure 2
doc.add_picture("paper/figures/fig2_priority_dist.png", width=Inches(4.5))
add_caption(doc, "Gambar 2. Distribusi Level Prioritas Dataset COBACEK")
doc.add_paragraph()

add_heading(doc, "3.2 Preprocessing dan Feature Extraction", level=2)
add_para(doc,
    "Deskripsi tiket melalui preprocessing standar: konversi ke lowercase, penghapusan whitespace "
    "berlebih, dan penghapusan karakter spesial. Feature extraction menggunakan TF-IDF:"
)
eq_p = doc.add_paragraph()
eq_r = eq_p.add_run("    x_i = TF-IDF(d_i, ngram=(1,2))                              (1)")
eq_r.italic = True
eq_r.font.size = Pt(11)
add_para(doc,
    "Di mana x_i adalah feature vector untuk tiket ke-i, dan ngram=(1,2) menggunakan unigram dan bigram. "
    "TF-IDF dipilih karena efektif untuk text classification dengan linear models dan memberikan "
    "interpretasi yang jelas tentang fitur-fitur penting."
)

add_heading(doc, "3.3 Experimental Setup", level=2)
add_para(doc,
    "Penelitian membandingkan empat pendekatan. BERT di-exclude dari experiment utama karena keterbatasan "
    "computational resources (GPU) dan fokus penelitian pada perbandingan classical ML vs Hybrid GenAI. "
    "Stratified 5-Fold Cross-Validation digunakan untuk evaluasi fair pada imbalanced dataset."
)

# Figure 3
doc.add_picture("paper/figures/fig3_pipeline.png", width=Inches(5.8))
add_caption(doc, "Gambar 3. Alur Metodologi Penelitian")
doc.add_paragraph()

add_heading(doc, "3.3.1 Konfigurasi Model", level=3)

add_heading(doc, "Support Vector Machine (SVM)", level=4)
add_para(doc,
    "SVM menggunakan Linear kernel melalui implementasi LinearSVC dari scikit-learn. "
    "Linear kernel dipilih karena TF-IDF features sudah berada di high-dimensional sparse space "
    "yang secara inheren linearly separable — kernel non-linear seperti RBF tidak memberikan "
    "keuntungan signifikan. Dua model SVM terpisah dilatih: satu untuk kategori (81 kelas) "
    "dan satu untuk prioritas (3 kelas). Pipeline: TF-IDF -> LinearSVC."
)

add_heading(doc, "Random Forest (RF)", level=4)
add_para(doc,
    "Random Forest dengan n_estimators=200, random_state=42. Dua model terpisah untuk kategori "
    "dan prioritas. RF dipilih sebagai representasi ensemble tree-based approach untuk "
    "dibandingkan dengan linear SVM."
)

add_heading(doc, "Logistic Regression (LR)", level=4)
add_para(doc,
    "Logistic Regression dengan max_iter=1000 sebagai probabilistic linear baseline. "
    "LR lebih sederhana dari SVM tapi powerful untuk text classification karena menghasilkan "
    "probability scores untuk setiap kelas."
)

add_heading(doc, "3.4 Evaluation Metrics", level=2)
add_para(doc, "Empat metrik digunakan untuk evaluasi komprehensif:")
for eq, desc in [
    ("Accuracy = (TP + TN) / (TP + TN + FP + FN)    (2)",
     "Proporsi prediksi benar dari total. Berguna tapi kurang sensitif terhadap class imbalance."),
    ("Macro Precision = (1/C) * SUM(TP_c / (TP_c + FP_c))    (3)",
     "Precision rata-rata per kelas dengan bobot equal. Adil untuk imbalanced dataset."),
    ("Macro Recall = (1/C) * SUM(TP_c / (TP_c + FN_c))    (4)",
     "Recall rata-rata per kelas dengan bobot equal."),
    ("Macro F1 = (1/C) * SUM(2*P_c*R_c / (P_c + R_c))    (5)",
     "Harmonic mean Precision dan Recall. Metrik utama penelitian ini."),
]:
    ep = doc.add_paragraph()
    ep.add_run(f"    {eq}").italic = True
    ep.runs[0].font.size = Pt(11)
    add_para(doc, desc, style="List Bullet")

add_heading(doc, "3.5 Hybrid SVM + GenAI Architecture", level=2)
add_para(doc,
    "Hybrid approach menggabungkan SVM baseline dengan GenAI correction (OpenAI gpt-5.4-mini) melalui "
    "empat tahap. Gambar 4 di bawah menggambarkan arsitektur lengkapnya:"
)

# Figure 4
doc.add_picture("paper/figures/fig4_hybrid_arch.png", width=Inches(4.8))
add_caption(doc, "Gambar 4. Arsitektur Pipeline Hybrid SVM + GenAI (4 tahap)")
doc.add_paragraph()

for stage, text in [
    ("Tahap 1 — SVM Baseline",
     "Model SVM Linear memprediksi kategori dan prioritas untuk seluruh 16.338 tiket."),
    ("Tahap 2 — Mismatch Detection",
     "Identifikasi tiket yang salah prediksi: M_i = {i | (cat_SVM != cat_true) OR (pri_SVM != pri_true)}. "
     "6.607 tiket (40.4%) dianggap mismatch dan menjadi kandidat koreksi GenAI."),
    ("Tahap 3 — GenAI Correction",
     "Untuk setiap tiket mismatch, OpenAI gpt-5.4-mini menerima deskripsi tiket dan prediksi SVM, "
     "kemudian mengembalikan kategori dan prioritas yang diperbaiki."),
    ("Tahap 4 — Result Combination",
     "Hasil akhir: {SVM correct rows} UNION {GenAI corrected rows}."),
]:
    p = doc.add_paragraph()
    r1 = p.add_run(stage + ": ")
    r1.bold = True
    r1.font.size = Pt(11)
    r2 = p.add_run(text)
    r2.font.size = Pt(11)

doc.add_page_break()

# ===========================================================
# BAB 4 — HASIL
# ===========================================================
add_heading(doc, "Bab 4. Hasil", level=1)

add_heading(doc, "4.1 Model Performance Comparison", level=2)
add_para(doc,
    "Tabel 1 merangkum hasil evaluasi dari dataset 16.338 tiket dengan 81 kategori dan 3 level prioritas. "
    "Semua model dievaluasi menggunakan Stratified 5-Fold Cross-Validation."
)

# Results table
col_heads = ["Model", "Task", "Accuracy", "Macro F1", "Weighted F1", "Macro Prec.", "Macro Recall", "Waktu (s)"]
result_rows = [
    ("SVM",                         "category", 0.8122, 0.2587, 0.8031, 0.3058, 0.2390,  67.82),
    ("SVM",                         "priority", 0.7250, 0.7114, 0.7228, 0.7311, 0.7008,  67.82),
    ("Random Forest",               "category", 0.7648, 0.1701, 0.7434, 0.2499, 0.1452, 306.98),
    ("Random Forest",               "priority", 0.7175, 0.6823, 0.7062, 0.7884, 0.6647, 306.98),
    ("Logistic Regression",         "category", 0.7734, 0.1093, 0.7428, 0.1379, 0.1047, 612.10),
    ("Logistic Regression",         "priority", 0.6422, 0.5962, 0.6274, 0.6711, 0.5880, 612.10),
    ("Hybrid SVM (gpt-5.4-mini)", "category", 0.6617, 0.2887, 0.7641, 0.3315, 0.4339, 9772.59),
    ("Hybrid SVM (gpt-5.4-mini)", "priority", 0.7383, 0.7226, 0.7376, 0.7250, 0.7206, 9772.59),
]

tbl2 = doc.add_table(rows=len(result_rows)+1, cols=len(col_heads))
tbl2.style = "Light Grid Accent 1"
for i, h in enumerate(col_heads):
    tbl2.rows[0].cells[i].text = h

for i, row in enumerate(result_rows, 1):
    tbl2.rows[i].cells[0].text = row[0]
    tbl2.rows[i].cells[1].text = row[1]
    for j in range(2, 7):
        tbl2.rows[i].cells[j].text = f"{row[j]:.4f}"
    tbl2.rows[i].cells[7].text = f"{row[7]:.2f}"

add_caption(doc, "Tabel 1. Perbandingan Performa Model pada Dataset COBACEK (n = 16.338)")
doc.add_paragraph()

# Figure 5 + 6 + 7
doc.add_picture("paper/figures/fig5_accuracy_compare.png", width=Inches(5.8))
add_caption(doc, "Gambar 5. Perbandingan Accuracy Antar Model untuk Kategori dan Prioritas")
doc.add_paragraph()

doc.add_picture("paper/figures/fig6_macro_f1_compare.png", width=Inches(5.8))
add_caption(doc, "Gambar 6. Perbandingan Macro F1 Antar Model (catatan: rendahnya Macro F1 kategori karena extreme class imbalance)")
doc.add_paragraph()

doc.add_picture("paper/figures/fig7_time_compare.png", width=Inches(5.0))
add_caption(doc, "Gambar 7. Perbandingan Waktu Komputasi Antar Model (skala logaritmik)")
doc.add_paragraph()

add_heading(doc, "4.2 Temuan Utama", level=2)

for title, body in [
    ("SVM: Best Performer untuk Kategori",
     "SVM Linear menunjukkan accuracy tertinggi untuk kategori (81.22%, Macro F1 = 0.2587). "
     "Untuk priority, SVM mencapai 72.50% accuracy (Macro F1 = 0.7114). Waktu komputasi efisien: 67.82 detik."),
    ("Random Forest dan Logistic Regression: Kompetitif tapi di bawah SVM",
     "RF mencapai accuracy kategori 76.48% dan LR 77.34% — keduanya di bawah SVM untuk kategori. "
     "LR khususnya lemah untuk priority (64.22%). Waktu komputasi RF (306.98s) dan LR (612.10s) "
     "jauh lebih lambat dari SVM."),
    ("Hybrid SVM + GenAI: Accuracy Paradox",
     "Temuan paling mengejutkan: Hybrid MENURUNKAN accuracy kategori dari 81.22% (SVM) menjadi "
     "66.17% — penurunan 14.05 percentage points. Sebaliknya, priority hanya naik 1.33 pp (72.50% -> 73.83%). "
     "Waktu komputasi 9.772 detik (~2.7 jam), yaitu 144x lebih lambat dari SVM standalone."),
    ("Macro F1 vs Accuracy: Diskrepansi Akibat Class Imbalance",
     "Macro F1 kategori sangat rendah (0.11-0.29) dibanding accuracy (0.66-0.81). Ini bukan tanda model buruk, "
     "melainkan cerminan extreme imbalance: 50+ kategori dengan < 30 sampel masing-masing menghasilkan "
     "per-class F1 yang sangat rendah untuk minority classes, menurunkan macro average secara drastis."),
]:
    p = doc.add_paragraph()
    r1 = p.add_run(title + ": ")
    r1.bold = True
    r1.font.size = Pt(11)
    p.add_run(body).font.size = Pt(11)

doc.add_page_break()

# ===========================================================
# BAB 5 — PEMBAHASAN
# ===========================================================
add_heading(doc, "Bab 5. Pembahasan", level=1)

add_heading(doc, "5.1 Mengapa SVM Linear Outperforms?", level=2)
add_para(doc,
    "SVM Linear mencapai accuracy tertinggi (81.22% kategori) karena beberapa faktor yang saling "
    "mendukung. TF-IDF menghasilkan high-dimensional sparse feature vectors yang secara inheren "
    "linearly separable — kondisi ideal untuk LinearSVC. Sementara model ensemble seperti RF "
    "memiliki overhead dari agregasi banyak decision trees pada sparse features, SVM langsung "
    "menemukan optimal hyperplane. Logistic Regression kalah karena regularisasi bawaan LR kurang "
    "optimal untuk imbalanced 81-class classification dibanding max-margin SVM."
)

add_heading(doc, "5.2 Accuracy Paradox: Mengapa GenAI Correction Menurunkan Akurasi?", level=2)
add_para(doc,
    "Temuan bahwa Hybrid SVM+GenAI menurunkan accuracy kategori dari 81.22% menjadi 66.17% adalah "
    "insight kritis. Beberapa faktor berkontribusi pada fenomena ini:"
)

for title, body in [
    ("Multi-label simultaneous correction",
     "GenAI diminta mengoreksi kategori DAN prioritas sekaligus. Bila tiket masuk mismatch karena "
     "hanya prioritas yang salah (kategori sudah benar), GenAI tetap mengubah KEDUA label untuk "
     "'konsistensi semantik'. Akibatnya, kategori yang awalnya benar menjadi salah."),
    ("Mismatch detection threshold terlalu broad",
     "Mismatch didefinisikan sebagai: salah kategori OR salah prioritas. Dari 6.607 mismatch rows, "
     "sebagian besar mungkin hanya mismatch di prioritas (kategori sudah benar). Koreksi GenAI pada "
     "kelompok ini merugikan accuracy kategori secara masif."),
    ("Limited context untuk GenAI",
     "GenAI hanya menerima deskripsi tiket dan prediksi SVM tanpa confidence scores. Tanpa visibility "
     "ke seberapa yakin SVM terhadap prediksinya, GenAI tidak tahu kapan harus mempertahankan vs "
     "mengoreksi prediksi yang sebenarnya sudah high-confidence benar."),
    ("GenAI hallucination pada domain-specific categories",
     "Dengan 81 kategori domain-spesifik, GenAI berpotensi menghasilkan kategori yang plausible "
     "secara linguistik tapi tidak sesuai dengan definisi internal helpdesk. SVM yang dilatih "
     "langsung pada data lebih akurat dalam memahami batas antar kategori spesifik ini."),
]:
    p = doc.add_paragraph()
    r1 = p.add_run(f"({title}): ")
    r1.bold = True
    r1.font.size = Pt(11)
    p.add_run(body).font.size = Pt(11)

add_heading(doc, "5.3 Implikasi untuk Production Deployment", level=2)

for title, body in [
    ("Hybrid tidak feasible untuk kategori",
     "Dengan accuracy turun 14 percentage points dan waktu 144x lebih lambat, Hybrid approach "
     "LEBIH BURUK dari SVM standalone untuk kategori classification. API cost dan latency "
     "tidak justified oleh degradasi accuracy ini."),
    ("SVM alone adalah pilihan optimal",
     "SVM Linear dengan TF-IDF adalah kombinasi terbaik: accuracy tertinggi, waktu tercepat, "
     "biaya komputasi minimal, dan tidak memerlukan API calls ke third-party service."),
    ("Hybrid mungkin berguna untuk priority saja",
     "Priority classification menunjukkan improvement kecil (+1.33 pp) dengan Hybrid, tapi "
     "cost-benefit masih questionable untuk improvement yang marginal."),
]:
    p = doc.add_paragraph()
    r1 = p.add_run(f"{title}: ")
    r1.bold = True
    r1.font.size = Pt(11)
    p.add_run(body).font.size = Pt(11)

add_heading(doc, "5.4 Rekomendasi Alternatif", level=2)
add_para(doc, "Berdasarkan analisis ini, alternatif yang lebih promising untuk future work:")
for item in [
    "Confidence-based routing: hanya apply GenAI untuk low-confidence SVM predictions (probability < 0.5), "
    "bukan blanket correction pada semua mismatches",
    "Separate task formulation: train GenAI khusus untuk kategori SAJA atau prioritas SAJA secara independen",
    "Ensemble voting: kombinasi SVM, RF, LR dengan majority voting untuk detect uncertain predictions",
    "GenAI sebagai feature engineer: gunakan GenAI untuk ekstrak structured features dari tiket, "
    "kemudian feed ke SVM — bukan sebagai direct classifier",
]:
    add_para(doc, item, style="List Bullet")

add_heading(doc, "5.5 Limitations", level=2)
add_para(doc, "Penelitian ini memiliki beberapa keterbatasan:")
for item in [
    "Dataset dari single organization — generalizability ke helpdesk lain belum terverifikasi",
    "BERT di-exclude karena keterbatasan GPU — perbandingan dengan fine-tuned transformer tidak dilakukan",
    "Hybrid GenAI correction menggunakan gpt-5.4-mini — model lain atau prompt engineering berbeda "
    "mungkin memberikan hasil berbeda",
    "Analisis error per-category tidak dilakukan secara mendalam",
]:
    add_para(doc, item, style="List Bullet")

doc.add_page_break()

# ===========================================================
# BAB 6 — KESIMPULAN
# ===========================================================
add_heading(doc, "Bab 6. Kesimpulan", level=1)

add_heading(doc, "6.1 Kesimpulan Utama", level=2)
for num, text in [
    ("1.", "SVM dengan Linear kernel adalah best performer untuk IT helpdesk ticket classification pada "
           "dataset ini: 81.22% accuracy kategori dan 72.50% priority, dengan waktu komputasi efisien (67.82s)."),
    ("2.", "Hybrid SVM+GenAI menunjukkan accuracy paradox: GenAI correction pada 6.607 tiket mismatch (40.4%) "
           "justru MENURUNKAN accuracy kategori sebesar 14.05 pp (81.22% -> 66.17%), sementara priority "
           "hanya naik 1.33 pp. Hybrid 144x lebih lambat dengan hasil lebih buruk untuk kategori."),
    ("3.", "Penyebab accuracy paradox: mismatch detection yang terlalu broad (salah satu label OR both), "
           "multi-label simultaneous correction, dan tidak adanya confidence-based filtering."),
    ("4.", "Macro F1 rendah untuk kategori (0.11-0.29) adalah artefak extreme class imbalance (81 kelas, "
           "many dengan < 5 sampel), bukan indikator kegagalan model."),
]:
    p = doc.add_paragraph()
    p.add_run(num + " ").bold = True
    p.add_run(text).font.size = Pt(11)
    p.runs[0].font.size = Pt(11)

add_heading(doc, "6.2 Kontribusi Penelitian", level=2)
for num, text in [
    ("1.", "Critical evaluation hybrid SVM+GenAI: discovery dan analisis accuracy paradox sebagai "
           "cautionary insight untuk research community."),
    ("2.", "Benchmark komprehensif 4 model pada real 16K-tiket helpdesk dataset — referensi valuable "
           "untuk practitioners yang membangun sistem helpdesk automation."),
    ("3.", "Analysis mendalam tentang mengapa simultaneous multi-label GenAI correction merusak accuracy, "
           "beserta rekomendasi concrete untuk future hybrid approaches."),
]:
    p = doc.add_paragraph()
    p.add_run(num + " ").bold = True
    p.add_run(text).font.size = Pt(11)
    p.runs[0].font.size = Pt(11)

add_heading(doc, "6.3 Rekomendasi Praktis", level=2)

for title, body in [
    ("Untuk accuracy optimal",
     "Deploy SVM Linear sebagai primary classifier. Accuracy 81.22% kategori dan 72.50% priority "
     "adalah solid baseline dengan minimal computational cost."),
    ("Untuk peningkatan tanpa GenAI",
     "Eksplorasi ensemble methods (voting SVM+RF+LR) atau confidence-based routing untuk defer "
     "uncertain predictions ke human review."),
    ("Jika ingin integrasi GenAI",
     "Gunakan GenAI untuk feature engineering atau selective correction (hanya low-confidence predictions), "
     "bukan blanket correction pada semua mismatches."),
]:
    p = doc.add_paragraph()
    r1 = p.add_run(title + ": ")
    r1.bold = True
    r1.font.size = Pt(11)
    p.add_run(body).font.size = Pt(11)

add_heading(doc, "6.4 Penutup", level=2)
add_para(doc,
    "Penelitian ini menunjukkan bahwa naive hybrid approach menggabungkan classical ML dengan GenAI tidak "
    "selalu memberikan improvement. Accuracy paradox yang ditemukan adalah pelajaran penting bahwa kombinasi "
    "model perlu didesain dengan careful consideration terhadap interaction effects dan task formulation. "
    "Untuk IT helpdesk ticket classification, SVM Linear alone memberikan performa excellent dengan "
    "computational cost minimal. Future research harus fokus pada selective correction strategies, "
    "confidence-based routing, dan separate task formulation untuk memanfaatkan GenAI secara efektif."
)

doc.add_page_break()

# ===========================================================
# DAFTAR PUSTAKA
# ===========================================================
add_heading(doc, "Daftar Pustaka", level=1)

refs = [
    "Devlin, J., Chang, M.-W., Lee, K., & Toutanova, K. (2019). BERT: Pre-training of Deep Bidirectional "
    "Transformers for Language Understanding. NAACL-HLT 2019.",
    "He, H., & Garcia, E. A. (2009). Learning from Imbalanced Data. IEEE Transactions on Knowledge and "
    "Data Engineering, 21(9), 1263–1284.",
    "Joachims, T. (1998). Text categorization with Support Vector Machines: Learning with many relevant "
    "features. ECML 1998.",
    "OpenAI. (2023). GPT-4 Technical Report. arXiv:2303.08774.",
    "Ratner, A., Bach, S. H., Ehrenberg, H., Fries, J., Wu, S., & Re, C. (2019). Snorkel: Rapid Training "
    "Data Creation with Weak Supervision. VLDB Endowment, 11(3), 269–282.",
    "Sebastiani, F. (2002). Machine learning in automated text categorization. ACM Computing Surveys, "
    "34(1), 1–47.",
    "Sun, C., Qiu, X., Xu, Y., & Huang, X. (2019). How to Fine-Tune BERT for Text Classification? "
    "CCL 2019.",
    "Yang, Y., & Liu, X. (1999). A re-examination of text categorization methods. SIGIR 1999.",
]
for i, ref in enumerate(refs, 1):
    p = doc.add_paragraph(style="List Number")
    p.add_run(ref).font.size = Pt(10.5)

# ===========================================================
# SAVE
# ===========================================================
out = "paper/IT_Helpdesk_Ticket_Classifier_Paper_V2.docx"
doc.save(out)
print(f"[OK] Paper saved: {out}")
print("Figures included: fig1-fig7 (category dist, priority dist, pipeline, hybrid arch, accuracy, macro f1, time)")
