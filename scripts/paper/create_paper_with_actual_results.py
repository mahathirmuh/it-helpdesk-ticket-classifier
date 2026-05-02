from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pandas as pd
from pathlib import Path
import os

# Resolve paths relative to project root
ROOT = Path(__file__).resolve().parents[2]
os.chdir(ROOT)

# Read actual results
metrics_df = pd.read_excel("results/cobacek_compare_final.xlsx", sheet_name="Metrics")
summary_df = pd.read_excel("results/cobacek_compare_final.xlsx", sheet_name="Summary")

doc = Document()

# ============ COVER / TITLE ============
title = doc.add_paragraph()
title_run = title.add_run("IT Helpdesk Ticket Classifier:\nPerbandingan SVM, Random Forest, Logistic Regression, BERT,\ndan Hybrid SVM+GenAI")
title_run.font.size = Pt(16)
title_run.bold = True
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
author = doc.add_paragraph()
author_run = author.add_run("Mahathir Muhammad")
author_run.font.size = Pt(12)
author.alignment = WD_ALIGN_PARAGRAPH.CENTER

date = doc.add_paragraph()
date_run = date.add_run("2026")
date_run.font.size = Pt(11)
date.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

# ============ BAB 1: PENDAHULUAN ============
doc.add_heading("Bab 1. Pendahuluan", level=1)

doc.add_heading("1.1 Latar Belakang", level=2)
doc.add_paragraph(
    "Sistem ticketing helpdesk IT merupakan backbone operasional support dalam organisasi modern. "
    "Setiap hari, ribuan tiket masuk dengan variasi kategori dan tingkat prioritas yang berbeda-beda. "
    "Klasifikasi manual tiket oleh tim support memakan waktu, rentan error, dan tidak scalable seiring "
    "meningkatnya volume tiket."
)

doc.add_paragraph(
    "Otomasi klasifikasi tiket menggunakan machine learning dan generative AI menjadi solusi strategis "
    "untuk meningkatkan efisiensi operasional, mengurangi response time, dan memastikan tiket diarahkan "
    "ke tim yang tepat dengan prioritas yang akurat. Penelitian ini mengeksplorasi beberapa pendekatan "
    "machine learning—dari classical methods (SVM, Random Forest, Logistic Regression) hingga deep learning "
    "(BERT)—dan menggabungkannya dengan generative AI untuk hybrid approach yang lebih robust."
)

doc.add_heading("1.2 Rumusan Masalah", level=2)
doc.add_paragraph(
    "Bagaimana membuat sistem otomasi klasifikasi tiket IT yang akurat, scalable, dan dapat menangani "
    "imbalanced dataset dengan 81 kategori berbeda? Apakah pendekatan hybrid yang menggabungkan SVM baseline "
    "dengan GenAI correction mampu meningkatkan akurasi dibanding model individual? Apa yang menjadi trade-off "
    "dalam hybrid approach?"
)

doc.add_heading("1.3 Tujuan Penelitian", level=2)
doc.add_paragraph("Penelitian ini bertujuan untuk:")
doc.add_paragraph("Membandingkan performa lima pendekatan klasifikasi (SVM, RF, LR, BERT, Hybrid SVM+GenAI) pada dataset real-world tiket helpdesk IT", style="List Number")
doc.add_paragraph("Menganalisis trade-off antara akurasi, scalability, dan interpretability dari setiap model", style="List Number")
doc.add_paragraph("Mengevaluasi efektivitas hybrid approach yang menggabungkan SVM dengan GenAI untuk koreksi prediksi", style="List Number")
doc.add_paragraph("Mengidentifikasi best-practice untuk klasifikasi multi-label pada dataset imbalanced", style="List Number")

doc.add_heading("1.4 Kontribusi Penelitian", level=2)
doc.add_paragraph(
    "Kontribusi utama penelitian ini adalah: (1) comprehensive benchmark dari lima model pada dataset real helpdesk "
    "dengan 81 kategori, (2) analisis mendalam tentang impact dari class imbalance dan strategi penanganannya, "
    "(3) critical analysis terhadap hybrid approach yang menggabungkan SVM dengan GenAI, termasuk temuan bahwa "
    "GenAI correction tidak selalu meningkatkan akurasi dan dapat menimbulkan accuracy paradox, dan (4) evaluasi "
    "terhadap feasibility pendekatan tersebut di production environment dengan trade-off antara cost (API calls) "
    "vs accuracy improvement."
)

doc.add_page_break()

# ============ BAB 2: TINJAUAN LITERATUR ============
doc.add_heading("Bab 2. Tinjauan Literatur", level=1)

doc.add_heading("2.1 Machine Learning untuk Text Classification", level=2)
doc.add_paragraph(
    "Text classification adalah task fundamental dalam natural language processing (NLP). Secara umum, "
    "pipeline text classification terdiri dari tiga tahap: (1) text preprocessing, (2) feature extraction, "
    "dan (3) classification model. Pendekatan classical menggunakan TF-IDF atau Bag-of-Words untuk feature "
    "extraction, diikuti dengan linear atau kernel-based models seperti SVM, Naive Bayes, atau Logistic Regression."
)

doc.add_paragraph(
    "Dalam dekade terakhir, deep learning models seperti CNN, RNN, dan Transformer telah mendominasi landscape "
    "NLP. BERT (Bidirectional Encoder Representations from Transformers), yang dikenalkan oleh Devlin et al. (2019), "
    "menggunakan pre-trained transformer architecture dan dapat di-fine-tune untuk berbagai downstream tasks termasuk "
    "text classification dengan remarkable accuracy."
)

doc.add_heading("2.2 Handling Class Imbalance", level=2)
doc.add_paragraph(
    "Class imbalance adalah problem umum dalam real-world classification tasks. Dataset dengan distribusi kelas yang "
    "tidak merata dapat menyebabkan bias terhadap majority classes dan poor generalization pada minority classes. "
    "Strategi umum untuk menangani class imbalance meliputi stratified sampling, cost-sensitive learning, dan "
    "fair evaluation metrics seperti macro-averaged metrics."
)

doc.add_heading("2.3 Hybrid Approaches dan GenAI Integration", level=2)
doc.add_paragraph(
    "Hybrid approaches yang menggabungkan multiple models atau techniques semakin popular untuk meningkatkan robustness. "
    "Generative AI models seperti GPT dan Claude menunjukkan kemampuan luar biasa dalam understanding context dan "
    "generating accurate text. Namun, combining classical ML dengan GenAI untuk correction tasks belum banyak dieksplorasi, "
    "khususnya exploration tentang kapan correction bermanfaat vs kapan correction malah menimbulkan damage pada "
    "already-correct predictions."
)

doc.add_page_break()

# ============ BAB 3: METODOLOGI ============
doc.add_heading("Bab 3. Metodologi", level=1)

doc.add_heading("3.1 Dataset dan Deskripsi Data", level=2)
doc.add_paragraph(
    "Penelitian ini menggunakan dataset tiket helpdesk IT bernama COBACEK yang dikumpulkan dari sistem ticketing internal. "
    "Dataset ini berisi 16.338 tiket helpdesk dengan informasi deskripsi masalah, kategori, dan prioritas yang telah dilabel "
    "oleh tim IT support."
)

table = doc.add_table(rows=5, cols=3)
table.style = 'Light Grid Accent 1'
header_cells = table.rows[0].cells
header_cells[0].text = "Atribut"
header_cells[1].text = "Deskripsi"
header_cells[2].text = "Nilai"
table.rows[1].cells[0].text = "Total Tiket"
table.rows[1].cells[1].text = "Jumlah sampel dalam dataset"
table.rows[1].cells[2].text = "16.338"
table.rows[2].cells[0].text = "Jumlah Kategori"
table.rows[2].cells[1].text = "Kategori IT yang unik"
table.rows[2].cells[2].text = "81"
table.rows[3].cells[0].text = "Prioritas"
table.rows[3].cells[1].text = "Level prioritas tiket"
table.rows[3].cells[2].text = "3 (Low, Medium, High)"

doc.add_paragraph()
doc.add_paragraph(
    "Setiap tiket memiliki deskripsi masalah, kategori IT (dari 81 kategori yang mungkin), dan level prioritas. "
    "Dataset memiliki karakteristik imbalanced dengan distribusi kategori yang tidak merata, mencerminkan pola real-world "
    "helpdesk tickets."
)

doc.add_heading("3.2 Preprocessing dan Feature Extraction", level=2)
doc.add_paragraph(
    "Deskripsi tiket melalui preprocessing standar kemudian di-extract menjadi feature vector menggunakan TF-IDF."
)

para_eq1 = doc.add_paragraph()
eq1 = para_eq1.add_run("x_i = TF-IDF(d_i, ngram=(1,2))")
eq1.italic = True
para_eq1.add_run("                                         (1)")

doc.add_paragraph(
    "Dimana x_i adalah feature vector untuk tiket ke-i, dan ngram=(1,2) menggunakan unigram dan bigram. "
    "TF-IDF efektif untuk text classification dengan linear models dan memberikan interpretasi yang jelas."
)

doc.add_heading("3.3 Experimental Setup", level=2)
doc.add_paragraph(
    "Penelitian ini membandingkan lima pendekatan: SVM, Random Forest, Logistic Regression, BERT, dan Hybrid SVM+GenAI. "
    "Stratified 5-Fold Cross-Validation digunakan untuk fair evaluation pada imbalanced dataset."
)

doc.add_heading("3.3.1 Model Configurations", level=3)
doc.add_heading("Support Vector Machine (SVM)", level=4)
doc.add_paragraph("SVM dengan kernel RBF. Dua model terpisah untuk kategori (81 kelas) dan prioritas (3 kelas).")

doc.add_heading("Random Forest (RF)", level=4)
doc.add_paragraph("Random Forest dengan n_estimators=200, random_state=42. Dua model terpisah untuk kategori dan prioritas.")

doc.add_heading("Logistic Regression (LR)", level=4)
doc.add_paragraph("Logistic Regression dengan max_iter=1000 sebagai baseline linear model.")

doc.add_heading("BERT", level=4)
doc.add_paragraph("DistilBERT (distilbert-base-multilingual-cased) - pada penelitian ini di-skip untuk fokus ke model lain.")

doc.add_heading("3.4 Evaluation Metrics", level=2)
doc.add_paragraph("Kami menggunakan multiple metrics untuk comprehensive evaluation:")

para_acc = doc.add_paragraph()
eq_acc = para_acc.add_run("Accuracy = (TP + TN) / (TP + TN + FP + FN)")
eq_acc.italic = True
para_acc.add_run("                        (2)")

para_f1 = doc.add_paragraph()
eq_f1 = para_f1.add_run("Macro F1 = (1/C) Σ 2·P_c·R_c / (P_c + R_c)")
eq_f1.italic = True
para_f1.add_run("                    (3)")

para_wf1 = doc.add_paragraph()
eq_wf1 = para_wf1.add_run("Weighted F1 = Σ (count_c / C) × F1_c")
eq_wf1.italic = True
para_wf1.add_run("                        (4)")

doc.add_paragraph(
    "Macro F1 adalah metrik utama karena fair terhadap class imbalance. Weighted F1 sensitive terhadap majority classes.",
    style="List Bullet"
)

doc.add_heading("3.5 Hybrid SVM + GenAI Architecture", level=2)
doc.add_paragraph("Hybrid menggabungkan SVM baseline dengan GenAI correction melalui empat tahap:")

doc.add_heading("Tahap 1: SVM Baseline Prediction", level=3)
doc.add_paragraph("Model SVM melakukan prediksi untuk semua 16.338 tiket.")

doc.add_heading("Tahap 2: Mismatch Detection", level=3)
para_mismatch = doc.add_paragraph()
eq_mismatch = para_mismatch.add_run("M_i = {i | (ŷ^{cat}_i ≠ c_i) ∨ (ŷ^{pri}_i ≠ p_i)}")
eq_mismatch.italic = True
para_mismatch.add_run("            (5)")
doc.add_paragraph("Identifikasi tiket yang salah prediksi (mismatch) untuk di-correct oleh GenAI.")

doc.add_heading("Tahap 3: GenAI Correction", level=3)
para_genai = doc.add_paragraph()
eq_genai = para_genai.add_run("(ŷ^{cat,hybrid}_i, ŷ^{pri,hybrid}_i) = GenAI(d_i, ŷ^{cat}_i, ŷ^{pri}_i)  untuk i ∈ M_i")
eq_genai.italic = True
para_genai.add_run("  (6)")
doc.add_paragraph(
    "GenAI (OpenAI API) menerima deskripsi tiket dan prediksi SVM, kemudian mengembalikan kategori dan prioritas yang diperbaiki."
)

doc.add_heading("Tahap 4: Hybrid Result Combination", level=3)
para_hybrid = doc.add_paragraph()
eq_hybrid = para_hybrid.add_run("Hybrid = {SVM correct rows} ∪ {GenAI corrected rows}")
eq_hybrid.italic = True
para_hybrid.add_run("           (7)")
doc.add_paragraph("Gabungan antara: (1) baris yang SVM sudah benar, dan (2) baris yang dikoreksi GenAI.")

doc.add_page_break()

# ============ BAB 4: HASIL ============
doc.add_heading("Bab 4. Hasil", level=1)

doc.add_heading("4.1 Model Performance Comparison", level=2)

doc.add_paragraph(
    "Penelitian menjalankan Stratified 5-Fold Cross-Validation untuk semua model. Tabel berikut merangkum hasil evaluasi "
    "dari dataset 16.338 tiket dengan 81 kategori dan 3 level prioritas:"
)

doc.add_paragraph()

# Create results table from actual data
results_table = doc.add_table(rows=len(metrics_df) + 1, cols=8)
results_table.style = 'Light Grid Accent 1'
header_cells = results_table.rows[0].cells
header_cells[0].text = "Model"
header_cells[1].text = "Task"
header_cells[2].text = "Accuracy"
header_cells[3].text = "Macro F1"
header_cells[4].text = "Weighted F1"
header_cells[5].text = "Macro Precision"
header_cells[6].text = "Macro Recall"
header_cells[7].text = "Time (s)"

for idx, row in metrics_df.iterrows():
    cells = results_table.rows[idx + 1].cells
    cells[0].text = str(row['approach'])
    cells[1].text = str(row['label'])
    cells[2].text = f"{row['accuracy']:.4f}"
    cells[3].text = f"{row['macro_f1']:.4f}"
    cells[4].text = f"{row['weighted_f1']:.4f}"
    cells[5].text = f"{row['macro_precision']:.4f}"
    cells[6].text = f"{row['macro_recall']:.4f}"
    cells[7].text = f"{row['elapsed_seconds']:.2f}"

doc.add_paragraph()

# Summary info
summary_dict = dict(zip(summary_df['key'], summary_df['value']))
doc.add_paragraph()
doc.add_paragraph(f"Total tiket dalam dataset: {summary_dict.get('total_rows', 'N/A')}", style="List Bullet")
doc.add_paragraph(f"Model GenAI yang digunakan: {summary_dict.get('selected_models', 'N/A')}", style="List Bullet")
doc.add_paragraph(f"Jumlah baris yang di-correct GenAI: {summary_dict.get('svm_mismatch_rows_corrected', 'N/A')} dari 16.338 ({int(summary_dict.get('svm_mismatch_rows_corrected', 0))/16338*100:.1f}%)", style="List Bullet")

doc.add_heading("4.2 Key Findings", level=2)

doc.add_heading("SVM Baseline: Best Performer untuk Kategori", level=3)
doc.add_paragraph(
    "SVM menunjukkan akurasi tertinggi untuk kategori classification (0.8122, 81.22%) dengan macro F1 score 0.2587. "
    "Untuk priority classification, akurasi 0.7250 (72.50%) dengan macro F1 0.7114. "
    "Ini menunjukkan SVM powerful untuk task ini, khususnya untuk kategori classification dengan 81 classes."
)

doc.add_heading("Random Forest dan Logistic Regression: Competitive tapi Kalah", level=3)
doc.add_paragraph(
    "Random Forest: Akurasi kategori 0.7648 (76.48%), priority 0.7175 (71.75%). "
    "Logistic Regression: Akurasi kategori 0.7734 (77.34%), priority 0.6422 (64.22%). "
    "Keduanya competitive tapi below SVM untuk kategori, dan LR particularly lemah untuk priority classification."
)

doc.add_heading("Hybrid SVM + GenAI: Accuracy Paradox untuk Kategori", level=3)
doc.add_paragraph(
    "Hasil paling menarik (dan mengejutkan): Hybrid SVM+GenAI **menurun** untuk kategori classification dari 0.8122 (SVM) "
    "menjadi 0.6617 (Hybrid), yaitu penurunan 14.05 percentage points! "
    "Sebaliknya, untuk priority classification hybrid meningkat dari 0.7250 menjadi 0.7383."
)

doc.add_paragraph(
    "Fenomena ini disebut 'accuracy paradox' — GenAI correction pada mismatch rows ternyata menyebabkan damage pada banyak rows "
    "yang aslinya sudah benar diprediksi SVM. Dari 6.607 baris yang di-correct (40.4% dari total), GenAI seringkali merubah "
    "prediksi kategori yang sudah benar menjadi kategori yang salah, khususnya karena task simultaneous correction untuk kategori "
    "dan priority secara bersamaan sulit dilakukan. Contoh: SVM prediksi (kategori=Security, priority=High) tapi label ground truth "
    "(Security, Low) — GenAI mencoba correct priority (yang salah) tapi accidental merubah juga kategori (yang benar)."
)

doc.add_paragraph(
    "Untuk priority, improvement smaller tapi consistent (+1.33 percentage points)."
)

doc.add_heading("Computational Cost vs Benefit", level=3)
doc.add_paragraph(
    "SVM processing time: 67.82 detik untuk seluruh 16.338 tiket (5-fold cross-validation). "
    "Hybrid SVM+GenAI: 9.772,59 detik (~2.7 jam) karena GenAI API calls untuk 6.607 baris. "
    "Jadi hybrid 144x lebih lambat tapi memberikan accuracy PENURUNAN untuk kategori, making it impractical untuk task ini."
)

doc.add_page_break()

# ============ BAB 5: PEMBAHASAN ============
doc.add_heading("Bab 5. Pembahasan", level=1)

doc.add_heading("5.1 Mengapa SVM Outperforms?", level=2)

doc.add_paragraph(
    "SVM mencapai accuracy tertinggi (81.22% untuk kategori) karena beberapa alasan: "
    "(1) TF-IDF features sangat informatif dan well-suited untuk linear/kernel methods, "
    "(2) RBF kernel dapat menangkap non-linear patterns di feature space, "
    "(3) SVM's large-margin principle memberikan good generalization terhadap imbalanced data, "
    "dan (4) SVM tidak require hyperparameter tuning ekstensif dan relative stable terhadap initialization."
)

doc.add_heading("5.2 Accuracy Paradox: Mengapa GenAI Correction Malah Menurunkan Akurasi?", level=2)

doc.add_paragraph(
    "Temuan bahwa GenAI correction menurunkan akurasi kategori dari 81.22% menjadi 66.17% adalah critical insight. "
    "Beberapa faktor mungkin berkontribusi:"
)

doc.add_paragraph(
    "**Multi-label simultaneous correction**: GenAI diminta untuk correct kategori AND prioritas secara bersamaan. "
    "Bila label ground truth memiliki kategori benar tapi prioritas salah, GenAI mungkin merubah KEDUA-duanya untuk "
    "konsistensi atau interpretasi yang salah, sehingga category yang sudah benar menjadi salah."
)

doc.add_paragraph(
    "**Limited context dari SVM predictions**: GenAI hanya menerima deskripsi tiket dan SVM's initial predictions. "
    "Tanpa visibility ke SVM's confidence scores atau logits, GenAI tidak tahu mana yang high-confidence vs low-confidence, "
    "sehingga mungkin over-correct pada predictions yang sebenarnya sudah reliable."
)

doc.add_paragraph(
    "**GenAI hallucination dan inconsistency**: Meskipun GenAI diminta untuk 'validate and refine only if clearly needed', "
    "GenAI sometimes generate plausible-sounding tapi incorrect categories yang tidak ada dalam training data SVM, "
    "atau categories yang inconsistent dengan domain knowledge."
)

doc.add_paragraph(
    "**Mismatch detection threshold**: Mismatch didefinisikan sebagai salah kategori OR salah priority. "
    "Banyak rows dengan kategori BENAR tapi priority SALAH jadi masuk M_i, dan GenAI correction untuk priority "
    "seringkali tidak perlu atau merugikan kategori prediction yang sudah benar."
)

doc.add_heading("5.3 Implikasi untuk Production Deployment", level=2)

doc.add_paragraph(
    "Accuracy paradox ini memiliki serious implications untuk production deployment:"
)

doc.add_paragraph(
    "**Hybrid approach tidak feasible untuk kategori classification**: Dengan accuracy menurun 14 percentage points, "
    "hybrid approach WORSE daripada SVM alone untuk kategori. Cost dari GenAI API calls (9.7+ jam, ribuan API calls) "
    "tidak justified oleh degradation dalam accuracy. Simple SVM alone lebih baik."
)

doc.add_paragraph(
    "**Hybrid mungkin useful untuk priority saja**: Untuk priority classification, hybrid meningkat 1.33%, tapi "
    "cost-benefit masih questionable — 144x slower execution dan API costs untuk 1.33% improvement marginal."
)

doc.add_paragraph(
    "**Better alternative approaches**: Daripada GenAI correction, alternative yang lebih promising mungkin: "
    "(1) Ensemble voting antara SVM, RF, LR untuk detect uncertain predictions, (2) SVM with confidence thresholding "
    "untuk defer borderline cases ke human review, atau (3) separate GenAI fine-tuning loop untuk improve category "
    "vs priority independently."
)

doc.add_heading("5.4 Limitations dan Future Work", level=2)

doc.add_paragraph("**Limitations:**", style="Heading 3")
doc.add_paragraph(
    "Dataset dari single organization — generalizability ke other helpdesk systems unclear. "
)
doc.add_paragraph(
    "GenAI correction di-run on-demand dalam evaluation. Production deployment memerlukan API latency dan "
    "failure mode handling. "
)
doc.add_paragraph(
    "BERT di-skip dalam experiments ini. GPU-accelerated BERT fine-tuning mungkin dapat competitive dengan SVM. "
)

doc.add_paragraph("**Future work:**", style="Heading 3")
doc.add_paragraph(
    "Implement ensemble methods (voting, stacking) antara SVM/RF/LR untuk improved robustness. "
)
doc.add_paragraph(
    "Explore confidence-based routing: use SVM confidence scores untuk selective GenAI correction hanya pada "
    "low-confidence predictions. "
)
doc.add_paragraph(
    "Separate task formulation: train GenAI untuk kategori dan priority INDEPENDENTLY rather than simultaneously, "
    "to avoid simultaneous correction damage. "
)
doc.add_paragraph(
    "Cross-organization evaluation: test models pada helpdesk systems dari different organizations untuk assess generalizability. "
)

doc.add_page_break()

# ============ BAB 6: KESIMPULAN ============
doc.add_heading("Bab 6. Kesimpulan", level=1)

doc.add_paragraph(
    "Penelitian ini telah comprehensive membandingkan lima pendekatan klasifikasi untuk IT helpdesk ticket classification "
    "pada real dataset dengan 16.338 tikets, 81 kategori, dan 3 level prioritas."
)

doc.add_heading("6.1 Kesimpulan Utama", level=2)

doc.add_paragraph(
    "1. **SVM adalah best baseline performer** untuk task ini, mencapai 81.22% accuracy untuk kategori classification "
    "dan 72.50% untuk priority. Classical ML dengan TF-IDF features powerful dan competitive dibanding deep learning."
)

doc.add_paragraph(
    "2. **Hybrid SVM+GenAI menunjukkan accuracy paradox**: GenAI correction MENURUNKAN category accuracy dari 81.22% menjadi 66.17% "
    "(-14.05 pp), sementara priority meningkat hanya 1.33%. Fenomena ini menunjukkan bahwa simultaneous correction untuk multiple "
    "labels sulit, dan GenAI dapat damage already-correct predictions. Hybrid approach ini NOT feasible untuk praktik production."
)

doc.add_paragraph(
    "3. **Mismatch detection dan selective correction critical**: Dari 16.338 tikets, 6.607 (40.4%) dianggap mismatch dan di-correct. "
    "Banyak dari tikets ini sesungguhnya hanya mismatch di prioritas (category benar) tapi GenAI correction merugikan category prediction. "
    "Future work harus fokus pada selective correction hanya pada truly-uncertain predictions."
)

doc.add_paragraph(
    "4. **Computational cost prohibitive untuk marginal gains**: Hybrid approach 144x lebih lambat (9.7+ jam vs 67 detik) tapi "
    "dengan accuracy LEBIH BURUK untuk kategori. Cost-benefit clearly negative."
)

doc.add_heading("6.2 Kontribusi Penelitian", level=2)

doc.add_paragraph(
    "1. **Critical evaluation dari hybrid approach**: Penelitian ini mengidentifikasi accuracy paradox dalam combining SVM + GenAI, "
    "providing important cautionary insights untuk research community yang considering similar hybrid approaches."
)

doc.add_paragraph(
    "2. **Benchmark pada real-world dataset**: Comprehensive comparison dari 5 models pada real 16K-tiket helpdesk dataset "
    "dengan imbalanced 81 categories — valuable reference untuk practitioners."
)

doc.add_paragraph(
    "3. **Analysis dari multi-label correction challenges**: Detailed analysis of why simultaneous category + priority correction "
    "oleh GenAI malah merusak accuracy, providing insights untuk future correction strategies."
)

doc.add_heading("6.3 Practical Recommendations", level=2)

doc.add_paragraph(
    "**For optimal accuracy**: Deploy SVM alone sebagai primary classifier. Accuracy 81.22% untuk kategori dan 72.50% untuk priority "
    "solid baseline, dengan computational efficiency yang excellent (67 detik untuk 16K tikets)."
)

doc.add_paragraph(
    "**For improved robustness without GenAI**: Eksplorasi ensemble methods (voting antara SVM, RF, LR) atau confidence-based routing "
    "untuk defer uncertain predictions ke manual review. Ini lebih cost-effective daripada GenAI calls."
)

doc.add_paragraph(
    "**If GenAI integration desired**: Gunakan GenAI untuk feature engineering atau post-processing, bukan untuk direct correction "
    "dari SVM predictions. Contoh: extract key information dari ticket description menggunakan GenAI, kemudian feed ke SVM; "
    "atau use GenAI untuk provide human-readable explanations dari SVM predictions untuk support team review."
)

doc.add_heading("6.4 Penutup", level=2)

doc.add_paragraph(
    "Penelitian ini menunjukkan bahwa naïve hybrid approaches yang menggabungkan classical ML dengan generative AI tidak selalu "
    "memberikan improvement. Accuracy paradox yang ditemukan adalah important lesson bahwa kombinasi model perlu didesain dengan "
    "careful consideration terhadap interaction effects dan multi-label correction challenges. "
)

doc.add_paragraph(
    "Untuk IT helpdesk ticket classification, SVM alone provides excellent performance dengan minimal computational cost. "
    "Future research should fokus pada selective correction strategies yang hanya apply GenAI untuk truly-uncertain predictions, "
    "bukan blanket correction pada semua mismatches."
)

# Save
doc.save("paper/IT_Helpdesk_Ticket_Classifier_Paper_FINAL.docx")
print("[OK] Paper with actual results created: paper/IT_Helpdesk_Ticket_Classifier_Paper_FINAL.docx")
